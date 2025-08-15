import streamlit as st
import pandas as pd
import tempfile
from datetime import datetime, date, time as dt_time
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from oauth2client.client import OAuth2Credentials
import httplib2
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt
import plotly.express as px
import re
import uuid
import time

# =============================================
# CONFIGURAÇÕES GERAIS
# =============================================
st.set_page_config(page_title="Timesheet Fiscal", layout="wide")
st.sidebar.markdown(f"📅 Hoje é: **{date.today().strftime('%d/%m/%Y')}**")

CSV_SEP = ";"
CSV_ENC = "utf-8-sig"
BASES = {
    "timesheet.csv": [
        "ID", "Usuário", "Nome", "Data", "Empresa", "Projeto", "Time",
        "Atividade", "Quantidade", "Horas Gastas", "Observações",
        "DataHoraLancamento"
    ],
    "empresas.csv": ["Codigo SAP", "Nome Empresa", "Descrição"],
    "projetos.csv": ["Nome Projeto", "Time", "Status"],
    "atividades.csv": ["Nome Atividade", "Projeto Vinculado", "Descrição", "Status"],
}

ADMIN_USERS = ["cvieira", "wreis", "waraujo", "iassis"]

# =============================================
# AUTENTICAÇÃO
# =============================================
@st.cache_data
def carregar_usuarios():
    usuarios_config = st.secrets.get("users", {})
    usuarios = {}
    for user, dados in usuarios_config.items():
        try:
            nome, senha = dados.split("|", 1)
            usuarios[user] = {"name": nome, "password": senha}
        except Exception:
            st.warning(f"Erro ao carregar usuário '{user}' nos secrets.")
    return usuarios

users = carregar_usuarios()

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = ""

if not st.session_state.logged_in:
    st.title("🔐 Login")
    username = st.text_input("Usuário")
    password = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        user = users.get(username)
        if user and user["password"] == password:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success(f"Bem-vindo, {user['name']}!")
            st.rerun()
        else:
            st.error("Usuário ou senha incorretos.")
    st.stop()

st.sidebar.image("PRIO_SEM_POLVO_PRIO_PANTONE_LOGOTIPO_Azul.png")
nome_usuario = users[st.session_state.username]["name"]
st.sidebar.success(f"Logado como: {nome_usuario}")
if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# =============================================
# GOOGLE DRIVE (PyDrive2) + UTILITÁRIOS
# =============================================

def _build_credentials_from_secrets():
    cred_dict = st.secrets["credentials"]
    credentials = OAuth2Credentials(
        access_token=cred_dict["access_token"],
        client_id=cred_dict["client_id"],
        client_secret=cred_dict["client_secret"],
        refresh_token=cred_dict["refresh_token"],
        token_expiry=datetime.strptime(cred_dict["token_expiry"], "%Y-%m-%dT%H:%M:%SZ"),
        token_uri=cred_dict["token_uri"],
        user_agent="streamlit-app/1.0",
        revoke_uri=cred_dict.get("revoke_uri")
    )
    http = httplib2.Http()
    credentials.refresh(http)
    return credentials

@st.cache_resource(show_spinner=False)
def conectar_drive():
    gauth = GoogleAuth()
    gauth.credentials = _build_credentials_from_secrets()
    drive = GoogleDrive(gauth)
    return drive

@st.cache_resource(show_spinner=False)
def obter_pasta_ts_fiscal_id():
    drive = conectar_drive()
    lista = drive.ListFile({
        'q': "title='ts-fiscal' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    }).GetList()
    if lista:
        return lista[0]['id']
    pasta = drive.CreateFile({
        'title': 'ts-fiscal',
        'mimeType': 'application/vnd.google-apps.folder'
    })
    pasta.Upload()
    return pasta['id']

# ---------- Locks ----------

def _locks_folder_id(drive, root_id):
    lista = drive.ListFile({'q': f"'{root_id}' in parents and title='locks' and mimeType='application/vnd.google-apps.folder' and trashed=false"}).GetList()
    if lista:
        return lista[0]['id']
    p = drive.CreateFile({'title': 'locks', 'mimeType': 'application/vnd.google-apps.folder', 'parents': [{'id': root_id}]})
    p.Upload()
    return p['id']

class DriveLock:
    def __init__(self, base_name: str, timeout_sec: int = 8):
        self.drive = conectar_drive()
        self.root_id = obter_pasta_ts_fiscal_id()
        self.locks_id = _locks_folder_id(self.drive, self.root_id)
        self.base_name = base_name
        self.lock_title = f"{base_name}.lock"
        self.timeout_sec = timeout_sec
        self.file = None

    def acquire(self):
        start = time.time()
        while True:
            existing = self.drive.ListFile({'q': f"'{self.locks_id}' in parents and title='{self.lock_title}' and trashed=false"}).GetList()
            if not existing:
                # create
                f = self.drive.CreateFile({'title': self.lock_title, 'parents': [{'id': self.locks_id}]})
                f.SetContentString(f"locked-by={st.session_state.username}; ts={datetime.now().isoformat()}")
                try:
                    f.Upload()
                    self.file = f
                    return True
                except Exception:
                    pass
            if time.time() - start > self.timeout_sec:
                return False
            time.sleep(0.4)

    def release(self):
        try:
            if self.file is not None:
                self.file.Delete()
        except Exception:
            pass

# ---------- Arquivos base ----------

def _ensure_base_exists(title: str):
    drive = conectar_drive()
    root_id = obter_pasta_ts_fiscal_id()
    files = drive.ListFile({'q': f"'{root_id}' in parents and title='{title}' and trashed=false"}).GetList()
    if files:
        return files[0]
    # criar novo com colunas padrão
    cols = BASES[title]
    df = pd.DataFrame(columns=cols)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv").name
    df.to_csv(tmp, sep=CSV_SEP, index=False, encoding=CSV_ENC)
    f = drive.CreateFile({'title': title, 'parents': [{'id': root_id}]})
    f.SetContentFile(tmp)
    f.Upload()
    return f


def _get_latest_by_title(title: str):
    drive = conectar_drive()
    root_id = obter_pasta_ts_fiscal_id()
    files = drive.ListFile({'q': f"'{root_id}' in parents and title='{title}' and trashed=false"}).GetList()
    if not files:
        return _ensure_base_exists(title)
    # se houver múltiplos, pega o mais recente por modifiedDate
    files.sort(key=lambda x: x.get('modifiedDate', ''), reverse=True)
    return files[0]


def _read_csv_file(file_obj):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv").name
    file_obj.GetContentFile(tmp)
    df = pd.read_csv(tmp, sep=CSV_SEP, encoding=CSV_ENC)
    meta = {
        'id': file_obj['id'],
        'title': file_obj['title'],
        'modifiedDate': file_obj.get('modifiedDate'),
        'version': file_obj.get('version'),
    }
    return df, meta


def carregar_base(title: str):
    f = _get_latest_by_title(title)
    return _read_csv_file(f)


def _save_csv_to_file(file_obj, df: pd.DataFrame):
    # normaliza Data (se existir) para ISO string
    if 'Data' in df.columns:
        df['Data'] = pd.to_datetime(df['Data'], errors='coerce').dt.strftime('%Y-%m-%d')
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv").name
    df.to_csv(tmp, sep=CSV_SEP, index=False, encoding=CSV_ENC)
    file_obj.SetContentFile(tmp)
    file_obj.Upload()
    return file_obj


def salvar_backup(df: pd.DataFrame, base_title: str, revision: str | None):
    drive = conectar_drive()
    root_id = obter_pasta_ts_fiscal_id()
    base_sem_ext = base_title.rsplit('.', 1)[0]
    pasta_nome = f"Backup_{base_sem_ext}"
    pasta_list = drive.ListFile({'q': f"'{root_id}' in parents and title='{pasta_nome}' and mimeType='application/vnd.google-apps.folder' and trashed=false"}).GetList()
    if pasta_list:
        backup_id = pasta_list[0]['id']
    else:
        p = drive.CreateFile({'title': pasta_nome, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [{'id': root_id}]})
        p.Upload()
        backup_id = p['id']
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    rev = f"rev-{revision}" if revision else "rev-unknown"
    fname = f"{base_sem_ext}__{ts}__{rev}.csv"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv").name
    df.to_csv(tmp, sep=CSV_SEP, index=False, encoding=CSV_ENC)
    arq = drive.CreateFile({'title': fname, 'parents': [{'id': backup_id}]})
    arq.SetContentFile(tmp)
    arq.Upload()

# ---------- Operações de escrita seguras ----------

def append_rows(title: str, df_new: pd.DataFrame):
    lock = DriveLock(title)
    if not lock.acquire():
        st.error("Sistema ocupado. Tente novamente em alguns segundos.")
        return False
    try:
        file = _get_latest_by_title(title)
        df_cur, meta = _read_csv_file(file)
        # alinhar colunas
        all_cols = sorted(set(df_cur.columns).union(df_new.columns))
        df_cur = df_cur.reindex(columns=all_cols)
        df_new = df_new.reindex(columns=all_cols)
        df_merged = pd.concat([df_cur, df_new], ignore_index=True)
        if 'ID' in df_merged.columns:
            df_merged = df_merged.drop_duplicates(subset=['ID'], keep='last')
        file = _save_csv_to_file(file, df_merged)
        salvar_backup(df_merged, title, file.get('version'))
        return True
    finally:
        lock.release()


def update_row_by_id(title: str, row_id: str, updates: dict):
    lock = DriveLock(title)
    if not lock.acquire():
        st.error("Sistema ocupado. Tente novamente em alguns segundos.")
        return False
    try:
        file = _get_latest_by_title(title)
        df, meta = _read_csv_file(file)
        if 'ID' not in df.columns:
            st.error("Base sem coluna ID. Não é possível editar com segurança.")
            return False
        mask = df['ID'] == row_id
        if not mask.any():
            st.error("Registro não encontrado. Recarregue a página.")
            return False
        for k, v in updates.items():
            if k in df.columns:
                df.loc[mask, k] = v
        file = _save_csv_to_file(file, df)
        salvar_backup(df, title, file.get('version'))
        return True
    finally:
        lock.release()


def delete_row_by_id(title: str, row_id: str):
    lock = DriveLock(title)
    if not lock.acquire():
        st.error("Sistema ocupado. Tente novamente em alguns segundos.")
        return False
    try:
        file = _get_latest_by_title(title)
        df, meta = _read_csv_file(file)
        if 'ID' not in df.columns:
            st.error("Base sem coluna ID. Não é possível excluir com segurança.")
            return False
        before = len(df)
        df = df[df['ID'] != row_id]
        after = len(df)
        if before == after:
            st.error("Registro não encontrado. Recarregue a página.")
            return False
        file = _save_csv_to_file(file, df)
        salvar_backup(df, title, file.get('version'))
        return True
    finally:
        lock.release()

# =============================================
# TRATAMENTO DE CAMPOS
# =============================================

def gerar_id_unico():
    return str(uuid.uuid4())


def formatar_horas(horas_input):
    if horas_input is None or str(horas_input).strip() == "":
        return None
    horas_input = str(horas_input).strip().replace(",", ".")
    pattern = re.fullmatch(r"(\d{1,2})[:;.,](\d{1,2})", horas_input)
    if pattern:
        h, m = map(int, pattern.groups())
        if 0 <= h < 24 and 0 <= m < 60:
            return f"{h:02d}:{m:02d}"
    try:
        decimal = float(horas_input)
        total_min = int(round(decimal * 60))
        h = total_min // 60
        m = total_min % 60
        return f"{h:02d}:{m:02d}"
    except Exception:
        return None


def normalizar_coluna_horas(df, coluna="Horas Gastas"):
    if coluna in df.columns:
        df[coluna] = df[coluna].astype(str).apply(formatar_horas)
    return df


def tratar_coluna_data(df, coluna="Data"):
    if coluna in df.columns:
        parsed = pd.to_datetime(df[coluna], errors="coerce", format="%Y-%m-%d")
        df["DataValida"] = parsed.notnull()
        df[coluna] = parsed
    return df

# =============================================
# MENU LATERAL
# =============================================

st.sidebar.title("📋 Menu")
menu = st.sidebar.radio("Navegar para:", [
    "🏠 Dashboard",
    "🏢 Cadastro de Empresas",
    "🗂️ Cadastro de Projetos e Atividades",
    "📝 Lançamento de Timesheet",
    "📄 Visualizar / Editar Timesheet",
    "📊 Avaliação de Performance — IA"
])

# =============================================
# CONTEÚDO: DASHBOARD
# =============================================

if menu == "🏠 Dashboard":
    st.title("📊 Painel de KPIs do Timesheet")
    df_timesheet, meta = carregar_base("timesheet.csv")
    df_timesheet = normalizar_coluna_horas(df_timesheet)
    df_timesheet = tratar_coluna_data(df_timesheet)

    if df_timesheet.empty:
        st.info("⚠️ Não há dados no timesheet para gerar dashboard.")
        st.stop()

    def converter_para_horas(horas_str):
        try:
            h, m = map(int, str(horas_str).strip().split(":"))
            return h + m / 60
        except Exception:
            return 0

    df_ts = df_timesheet[df_timesheet.get("DataValida", True) == True].copy()
    df_ts["Horas"] = df_ts["Horas Gastas"].apply(converter_para_horas)

    # Filtros
    st.sidebar.subheader("🔍 Filtros")
    df_ts["Data"] = pd.to_datetime(df_ts["Data"], errors="coerce")
    periodo_padrao = [
        (df_ts["Data"].min().date() if not df_ts.empty else date.today()),
        (df_ts["Data"].max().date() if not df_ts.empty else date.today())
    ]
    data_inicial, data_final = st.sidebar.date_input("Período:", periodo_padrao)

    empresa = st.sidebar.selectbox("Empresa:", ["Todas"] + sorted(df_ts["Empresa"].dropna().unique().tolist()))
    projeto = st.sidebar.selectbox("Projeto:", ["Todos"] + sorted(df_ts["Projeto"].dropna().unique().tolist()))
    squad = st.sidebar.selectbox("Time:", ["Todos"] + sorted(df_ts["Time"].dropna().unique().tolist()))
    atividade = st.sidebar.selectbox("Atividade:", ["Todas"] + sorted(df_ts["Atividade"].dropna().unique().tolist()))

    df_filtrado = df_ts[(df_ts["Data"].dt.date >= data_inicial) & (df_ts["Data"].dt.date <= data_final)].copy()
    if empresa != "Todas":
        df_filtrado = df_filtrado[df_filtrado["Empresa"] == empresa]
    if projeto != "Todos":
        df_filtrado = df_filtrado[df_filtrado["Projeto"] == projeto]
    if squad != "Todos":
        df_filtrado = df_filtrado[df_filtrado["Time"] == squad]
    if atividade != "Todas":
        df_filtrado = df_filtrado[df_filtrado["Atividade"] == atividade]

    total_horas = df_filtrado["Horas"].sum()
    total_registros = len(df_filtrado)
    total_colaboradores = df_filtrado["Nome"].nunique()
    total_projetos = df_filtrado["Projeto"].nunique()

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("⏰ Total de Horas", f"{total_horas:.2f}")
    c2.metric("📄 Total Registros", total_registros)
    c3.metric("👤 Colaboradores", total_colaboradores)
    c4.metric("🏗️ Projetos", total_projetos)

    if not df_filtrado.empty:
        st.subheader("🏗️ Horas por Projeto")
        gp = df_filtrado.groupby("Projeto")["Horas"].sum().reset_index().sort_values(by="Horas", ascending=False)
        st.plotly_chart(px.bar(gp, x="Projeto", y="Horas", text_auto='.2s'), use_container_width=True)

        st.subheader("🚀 Horas por Time")
        gt = df_filtrado.groupby("Time")["Horas"].sum().reset_index().sort_values(by="Horas", ascending=False)
        st.plotly_chart(px.bar(gt, x="Time", y="Horas", text_auto='.2s'), use_container_width=True)

        st.subheader("🗒️ Horas por Atividade")
        ga = df_filtrado.groupby("Atividade")["Horas"].sum().reset_index().sort_values(by="Horas", ascending=False)
        st.plotly_chart(px.bar(ga.head(), x="Atividade", y="Horas", text_auto='.2s'), use_container_width=True)

        st.subheader("🏢 Horas por Empresa")
        ge = df_filtrado.groupby("Empresa")["Horas"].sum().reset_index().sort_values(by="Horas", ascending=False)
        st.plotly_chart(px.pie(ge, names="Empresa", values="Horas", hole=0.4), use_container_width=True)

        st.subheader("👤 Horas por Colaborador")
        gc = df_filtrado.groupby("Nome")["Horas"].sum().reset_index().sort_values(by="Horas", ascending=False)
        st.plotly_chart(px.bar(gc, x="Nome", y="Horas", text_auto='.2s'), use_container_width=True)

        st.subheader("📅 Evolução de Horas no Tempo (Por Dia)")
        gt = (
            df_filtrado.assign(Dia=df_filtrado["Data"].dt.date)
            .groupby("Dia", as_index=False)["Horas"].sum()
            .sort_values("Dia")
        )
        fig = px.line(gt, x="Dia", y="Horas", markers=True)
        fig.update_xaxes(title="Dia", type="category")
        fig.update_yaxes(title="Horas")
        st.plotly_chart(fig, use_container_width=True)

# =============================================
# CONTEÚDO: EMPRESAS
# =============================================

elif menu == "🏢 Cadastro de Empresas":
    st.title("🏢 Cadastro de Empresas (Códigos SAP)")
    st.subheader("📥 Inserir nova empresa")

    df_empresas, _ = carregar_base("empresas.csv")

    with st.form("form_empresa"):
        c1, c2 = st.columns([2, 4])
        with c1:
            codigo = st.text_input("Código SAP")
        with c2:
            nome = st.text_input("Nome da Empresa")
        descricao = st.text_area("Descrição (opcional)", height=100)
        if st.form_submit_button("💾 Salvar Empresa"):
            if not codigo or not nome:
                st.warning("⚠️ Código SAP e Nome são obrigatórios.")
            else:
                if "Codigo SAP" in df_empresas.columns and codigo in df_empresas["Codigo SAP"].astype(str).values:
                    st.warning("⚠️ Já existe uma empresa cadastrada com este Código SAP.")
                else:
                    nova = pd.DataFrame({
                        "Codigo SAP": [str(codigo).strip()],
                        "Nome Empresa": [str(nome).strip()],
                        "Descrição": [str(descricao).strip()]
                    })
                    if append_rows("empresas.csv", nova):
                        st.success("✅ Empresa cadastrada com sucesso!")

    st.markdown("---")
    st.markdown("### 🏢 Empresas Cadastradas")
    df_empresas, _ = carregar_base("empresas.csv")
    st.dataframe(df_empresas, use_container_width=True)

    st.markdown("---")
    st.markdown("### 🛠️ Editar ou Excluir Empresa")
    if not df_empresas.empty and "Codigo SAP" in df_empresas.columns:
        empresa_sel = st.selectbox("Selecione a empresa pelo Código SAP:", df_empresas["Codigo SAP"].astype(str).tolist())
        registro = df_empresas[df_empresas["Codigo SAP"].astype(str) == str(empresa_sel)]
        if not registro.empty:
            row = registro.iloc[0]
            novo_nome = st.text_input("Novo Nome da Empresa", value=row.get("Nome Empresa", ""))
            nova_desc = st.text_area("Nova Descrição", value=row.get("Descrição", ""))
            cols = st.columns(2)
            with cols[0]:
                if st.button("✏️ Atualizar Empresa"):
                    idx_mask = (df_empresas["Codigo SAP"].astype(str) == str(empresa_sel))
                    if idx_mask.any():
                        # atualizar por regravação completa (pequena base) — mas com lock e backup
                        df_empresas.loc[idx_mask, "Nome Empresa"] = novo_nome.strip()
                        df_empresas.loc[idx_mask, "Descrição"] = nova_desc.strip()
                        # usa update_row_by_id? aqui não há ID; salvamos reescrevendo base inteira com lock
                        if delete_row_by_id("empresas.csv", row.get("ID", "__noid__")):
                            pass
                        # Como empresas não tem ID no seu legado, fazemos overwrite seguro com append_rows refazendo a base
                        # Estratégia simples: montar base inteira e salvar via _save_csv_to_file
                        lock = DriveLock("empresas.csv")
                        if lock.acquire():
                            try:
                                file = _get_latest_by_title("empresas.csv")
                                _save_csv_to_file(file, df_empresas)
                                salvar_backup(df_empresas, "empresas.csv", file.get('version'))
                                st.success("✅ Empresa atualizada!")
                                st.rerun()
                            finally:
                                lock.release()
            with cols[1]:
                if st.button("🗑️ Excluir Empresa"):
                    confirmar = st.radio("⚠️ Confirmar exclusão?", ["Não", "Sim"], horizontal=True, key="conf_emp")
                    if confirmar == "Sim":
                        df_empresas = df_empresas[df_empresas["Codigo SAP"].astype(str) != str(empresa_sel)]
                        lock = DriveLock("empresas.csv")
                        if lock.acquire():
                            try:
                                file = _get_latest_by_title("empresas.csv")
                                _save_csv_to_file(file, df_empresas)
                                salvar_backup(df_empresas, "empresas.csv", file.get('version'))
                                st.success("✅ Empresa excluída!")
                                st.rerun()
                            finally:
                                lock.release()
    else:
        st.info("🚩 Nenhuma empresa cadastrada até o momento.")

# =============================================
# CONTEÚDO: PROJETOS & ATIVIDADES
# =============================================

elif menu == "🗂️ Cadastro de Projetos e Atividades":
    st.title("🗂️ Cadastro de Projetos e Atividades")
    st.markdown("## 🏗️ Projetos")

    df_projetos, _ = carregar_base("projetos.csv")

    with st.form("form_projeto"):
        nome_projeto = st.text_input("Nome do Projeto")
        desc_time = st.selectbox("Time", ["Ambos", "Diretos", "Indiretos"])
        status_projeto = st.selectbox("Status do Projeto", ["Não Iniciado", "Em Andamento", "Concluído"])
        if st.form_submit_button("💾 Salvar Projeto"):
            if not nome_projeto:
                st.warning("⚠️ O nome do projeto é obrigatório.")
            elif "Nome Projeto" in df_projetos.columns and nome_projeto in df_projetos["Nome Projeto"].astype(str).values:
                st.warning("⚠️ Já existe um projeto com este nome.")
            else:
                novo = pd.DataFrame({
                    "Nome Projeto": [nome_projeto.strip()],
                    "Time": [desc_time.strip()],
                    "Status": [status_projeto]
                })
                if append_rows("projetos.csv", novo):
                    st.success("✅ Projeto cadastrado com sucesso!")

    df_projetos, _ = carregar_base("projetos.csv")
    st.dataframe(df_projetos, use_container_width=True)

    st.markdown("### 🔧 Editar ou Excluir Projeto")
    if not df_projetos.empty and "Nome Projeto" in df_projetos.columns:
        projeto_sel = st.selectbox("Selecione o Projeto:", df_projetos["Nome Projeto"].astype(str).tolist())
        idx = df_projetos[df_projetos["Nome Projeto"].astype(str) == str(projeto_sel)].index
        if not idx.empty:
            row = df_projetos.loc[idx[0]]
            novo_nome = st.text_input("Novo Nome do Projeto", value=row.get("Nome Projeto", ""))
            nova_desc = st.selectbox("Alterar Time", ["Ambos", "Diretos", "Indiretos"], index=["Ambos", "Diretos", "Indiretos"].index(row.get("Time", "Ambos")))
            novo_status = st.selectbox("Novo Status", ["Não Iniciado", "Em Andamento", "Concluído"], index=["Não Iniciado", "Em Andamento", "Concluído"].index(row.get("Status", "Não Iniciado")))
            c1, c2 = st.columns(2)
            with c1:
                if st.button("✏️ Atualizar Projeto"):
                    df_projetos.loc[idx, "Nome Projeto"] = novo_nome.strip()
                    df_projetos.loc[idx, "Time"] = nova_desc.strip()
                    df_projetos.loc[idx, "Status"] = novo_status
                    lock = DriveLock("projetos.csv")
                    if lock.acquire():
                        try:
                            file = _get_latest_by_title("projetos.csv")
                            _save_csv_to_file(file, df_projetos)
                            salvar_backup(df_projetos, "projetos.csv", file.get('version'))
                            st.success("✅ Projeto atualizado!")
                            st.rerun()
                        finally:
                            lock.release()
            with c2:
                if st.button("🗑️ Excluir Projeto"):
                    confirmar = st.radio("⚠️ Confirmar Exclusão?", ["Não", "Sim"], horizontal=True)
                    if confirmar == "Sim":
                        df_projetos = df_projetos.drop(idx)
                        lock = DriveLock("projetos.csv")
                        if lock.acquire():
                            try:
                                file = _get_latest_by_title("projetos.csv")
                                _save_csv_to_file(file, df_projetos)
                                salvar_backup(df_projetos, "projetos.csv", file.get('version'))
                                st.success("✅ Projeto excluído!")
                                st.rerun()
                            finally:
                                lock.release()

    # ATIVIDADES
    st.markdown("---")
    st.markdown("## 🗒️ Atividades")
    df_atividades, _ = carregar_base("atividades.csv")

    with st.form("form_atividade"):
        nome_atividade = st.text_input("Nome da Atividade")
        projeto_vinc = st.selectbox("Projeto Vinculado", df_projetos["Nome Projeto"].astype(str).tolist() if not df_projetos.empty else [])
        descricao_atv = st.text_area("Descrição da Atividade")
        status_atv = st.selectbox("Status da Atividade", ["Não Iniciada", "Em Andamento", "Concluída"])
        if st.form_submit_button("💾 Salvar Atividade"):
            if not nome_atividade:
                st.warning("⚠️ O nome da atividade é obrigatório.")
            elif "Nome Atividade" in df_atividades.columns and nome_atividade in df_atividades["Nome Atividade"].astype(str).values:
                st.warning("⚠️ Já existe uma atividade com este nome.")
            else:
                novo = pd.DataFrame({
                    "Nome Atividade": [nome_atividade.strip()],
                    "Projeto Vinculado": [projeto_vinc.strip()],
                    "Descrição": [descricao_atv.strip()],
                    "Status": [status_atv]
                })
                if append_rows("atividades.csv", novo):
                    st.success("✅ Atividade cadastrada com sucesso!")

    df_atividades, _ = carregar_base("atividades.csv")
    st.dataframe(df_atividades, use_container_width=True)

    st.markdown("### 🔧 Editar ou Excluir Atividade")
    if not df_atividades.empty and "Nome Atividade" in df_atividades.columns:
        atv_sel = st.selectbox("Selecione a Atividade:", df_atividades["Nome Atividade"].astype(str).tolist())
        idx = df_atividades[df_atividades["Nome Atividade"].astype(str) == str(atv_sel)].index
        if not idx.empty:
            row = df_atividades.loc[idx[0]]
            novo_nome = st.text_input("Novo Nome da Atividade", value=row.get("Nome Atividade", ""))
            novo_proj = st.selectbox("Novo Projeto Vinculado", df_projetos["Nome Projeto"].astype(str).tolist(), index=max(0, df_projetos["Nome Projeto"].astype(str).tolist().index(row.get("Projeto Vinculado", df_projetos["Nome Projeto"].astype(str).tolist()[0]))))
            nova_desc = st.text_area("Nova Descrição", value=row.get("Descrição", ""))
            novo_status = st.selectbox("Novo Status", ["Não Iniciada", "Em Andamento", "Concluída"], index=["Não Iniciada", "Em Andamento", "Concluída"].index(row.get("Status", "Não Iniciada")))
            c1, c2 = st.columns(2)
            with c1:
                if st.button("✏️ Atualizar Atividade"):
                    df_atividades.loc[idx, "Nome Atividade"] = novo_nome.strip()
                    df_atividades.loc[idx, "Projeto Vinculado"] = novo_proj.strip()
                    df_atividades.loc[idx, "Descrição"] = nova_desc.strip()
                    df_atividades.loc[idx, "Status"] = novo_status
                    lock = DriveLock("atividades.csv")
                    if lock.acquire():
                        try:
                            file = _get_latest_by_title("atividades.csv")
                            _save_csv_to_file(file, df_atividades)
                            salvar_backup(df_atividades, "atividades.csv", file.get('version'))
                            st.success("✅ Atividade atualizada!")
                            st.rerun()
                        finally:
                            lock.release()
            with c2:
                if st.button("🗑️ Excluir Atividade"):
                    confirmar = st.radio("⚠️ Confirmar Exclusão?", ["Não", "Sim"], horizontal=True)
                    if confirmar == "Sim":
                        df_atividades = df_atividades.drop(idx)
                        lock = DriveLock("atividades.csv")
                        if lock.acquire():
                            try:
                                file = _get_latest_by_title("atividades.csv")
                                _save_csv_to_file(file, df_atividades)
                                salvar_backup(df_atividades, "atividades.csv", file.get('version'))
                                st.success("✅ Atividade excluída!")
                                st.rerun()
                            finally:
                                lock.release()

# =============================================
# CONTEÚDO: LANÇAMENTO TS
# =============================================

elif menu == "📝 Lançamento de Timesheet":
    st.title("📝 Lançamento de Timesheet")
    st.subheader("⏱️ Registro de Horas")

    usuario_logado = st.session_state.username
    nome_usuario = users[usuario_logado]["name"]

    df_empresas, _ = carregar_base("empresas.csv")
    df_projetos, _ = carregar_base("projetos.csv")
    df_atividades, _ = carregar_base("atividades.csv")

    projeto = st.selectbox(
        "Projeto",
        sorted(df_projetos["Nome Projeto"].dropna().unique()) if not df_projetos.empty else ["Sem projetos cadastrados"]
    )

    df_atividades_filtrado = df_atividades[df_atividades["Projeto Vinculado"].astype(str) == str(projeto)]
    atividade = st.selectbox(
        "Atividade",
        sorted(df_atividades_filtrado["Nome Atividade"].dropna().unique()) if not df_atividades_filtrado.empty else ["Sem atividades para este projeto"]
    )

    squad = st.selectbox(
        "Time",
        sorted(df_projetos[df_projetos["Nome Projeto"].astype(str) == str(projeto)]["Time"].dropna().unique()) if not df_projetos.empty else ["Sem projetos cadastrados"]
    )

    with st.form("form_timesheet"):
        data_sel = st.date_input("Data", value=date.today())
        empresa = st.selectbox(
            "Empresa (Código SAP)",
            sorted(df_empresas["Codigo SAP"].dropna().astype(str).unique()) if not df_empresas.empty else ["Sem empresas cadastradas"]
        )
        quantidade = st.number_input("Quantidade Tarefas", min_value=0, step=1)
        tempo = st.time_input("Horas Gastas", value=dt_time(0, 0))
        horas = f"{tempo.hour:02d}:{tempo.minute:02d}"
        observacoes = st.text_area("Observações", placeholder="Descreva detalhes relevantes sobre este lançamento...", height=120, max_chars=500)
        if st.form_submit_button("💾 Registrar"):
            if horas == "00:00":
                st.warning("⚠️ O campo Horas Gastas não pode ser 00:00.")
            elif not projeto or not atividade or not empresa:
                st.warning("⚠️ Preencha todos os campos obrigatórios antes de registrar.")
            else:
                id_registro = gerar_id_unico()
                datahora_lanc = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                novo = pd.DataFrame({
                    "ID": [id_registro],
                    "Usuário": [usuario_logado],
                    "Nome": [nome_usuario],
                    "Data": [data_sel.strftime('%Y-%m-%d')],
                    "Empresa": [str(empresa)],
                    "Projeto": [str(projeto)],
                    "Time": [str(squad)],
                    "Atividade": [str(atividade)],
                    "Quantidade": [int(quantidade)],
                    "Horas Gastas": [horas],
                    "Observações": [observacoes.replace('\n', ' ').replace(';', ',').strip()],
                    "DataHoraLancamento": [datahora_lanc]
                })
                if append_rows("timesheet.csv", novo):
                    st.success("✅ Registro salvo no Timesheet com sucesso!")

# =============================================
# CONTEÚDO: VISUALIZAR / EDITAR TS
# =============================================

elif menu == "📄 Visualizar / Editar Timesheet":
    st.title("📄 Visualizar, Editar ou Excluir Timesheet")
    usuario_logado = st.session_state.username

    df_ts, _ = carregar_base("timesheet.csv")
    df_ts = normalizar_coluna_horas(df_ts)
    df_ts = tratar_coluna_data(df_ts)

    if usuario_logado not in ADMIN_USERS:
        df_ts = df_ts[df_ts["Usuário"] == usuario_logado]

    st.sidebar.subheader("🔍 Filtros")
    if df_ts.empty:
        period = [date.today(), date.today()]
    else:
        period = [
            (df_ts["Data"].min().date() if pd.notnull(df_ts["Data"].min()) else date.today()),
            (df_ts["Data"].max().date() if pd.notnull(df_ts["Data"].max()) else date.today())
        ]
    data_inicial, data_final = st.sidebar.date_input("Período:", period)

    empresa = st.sidebar.selectbox("Empresa:", ["Todas"] + sorted(df_ts["Empresa"].dropna().unique().tolist()) if not df_ts.empty else ["Todas"])
    projeto = st.sidebar.selectbox("Projeto:", ["Todos"] + sorted(df_ts["Projeto"].dropna().unique().tolist()) if not df_ts.empty else ["Todos"])
    squad = st.sidebar.selectbox("Time:", ["Todos"] + sorted(df_ts["Time"].dropna().unique().tolist()) if not df_ts.empty else ["Todos"])
    atividade = st.sidebar.selectbox("Atividade:", ["Todas"] + sorted(df_ts["Atividade"].dropna().unique().tolist()) if not df_ts.empty else ["Todas"])

    if usuario_logado in ADMIN_USERS:
        usuario_sel = st.sidebar.selectbox("Nome:", ["Todos"] + sorted(df_ts["Nome"].dropna().unique().tolist()) if not df_ts.empty else ["Todos"])
    else:
        usuario_sel = usuario_logado

    df_f = df_ts.copy()
    if empresa != "Todas":
        df_f = df_f[df_f["Empresa"] == empresa]
    if projeto != "Todos":
        df_f = df_f[df_f["Projeto"] == projeto]
    if squad != "Todos":
        df_f = df_f[df_f["Time"] == squad]
    if atividade != "Todas":
        df_f = df_f[df_f["Atividade"] == atividade]
    if usuario_sel != "Todos":
        df_f = df_f[df_f["Usuário"] == usuario_sel]

    df_f = df_f[(df_f["Data"].dt.date >= data_inicial) & (df_f["Data"].dt.date <= data_final)].sort_values(by="Data")

    df_visual = df_f.copy()
    df_visual["Data"] = df_visual["Data"].dt.strftime("%d/%m/%Y")
    df_visual = df_visual.rename(columns={"DataHoraLancamento": "Data de Registro"})

    cols_ordem = [c for c in df_visual.columns if c not in ["ID", "Data de Registro"]] + ["Data de Registro", "ID"]
    df_visual = df_visual[cols_ordem]

    st.markdown(f"### 🔍 {len(df_visual)} registros encontrados")
    if df_visual.empty:
        st.info("🚩 Nenhum registro encontrado com os filtros aplicados.")
        st.stop()
    else:
        st.dataframe(df_visual, use_container_width=True)

    st.markdown("---")
    st.subheader("✏️ Editar um Registro")
    indice = st.selectbox("Selecione o índice para editar:", df_f.index.tolist())
    linha = df_f.loc[indice]

    col_editar = st.selectbox("Coluna:", [
        "Data", "Nome", "Empresa", "Projeto", "Atividade", "Quantidade", "Horas Gastas", "Observações"
    ])

    valor_atual = linha[col_editar]
    if col_editar == "Data":
        novo_valor = st.date_input("Nova Data", value=valor_atual.date() if pd.notnull(valor_atual) else date.today())
        novo_valor = pd.to_datetime(novo_valor).strftime('%Y-%m-%d')
    elif col_editar == "Quantidade":
        novo_valor = st.number_input("Nova Quantidade", value=int(valor_atual) if pd.notnull(valor_atual) else 0)
    else:
        novo_valor = st.text_input("Novo Valor", value=str(valor_atual) if pd.notnull(valor_atual) else "").replace('\n', ' ').replace(';', ',').strip()

    if st.button("💾 Atualizar Registro"):
        id_editar = linha.get("ID", "")
        if not id_editar:
            st.error("❌ Este registro não possui ID. Não é possível editar com segurança.")
        else:
            ok = update_row_by_id("timesheet.csv", id_editar, {col_editar: novo_valor})
            if ok:
                st.success("✅ Registro atualizado com sucesso!")
                st.rerun()

    st.markdown("---")
    st.subheader("🗑️ Excluir um Registro")
    indice_excluir = st.selectbox("Índice para excluir:", df_f.index.tolist(), key="excluir")
    linha_x = df_f.loc[indice_excluir]
    st.markdown("**Registro selecionado:**")
    st.json({k: (v.strftime('%Y-%m-%d') if isinstance(v, (pd.Timestamp,)) else (v if pd.notnull(v) else None)) for k, v in linha_x.to_dict().items()})

    confirmar = st.radio("⚠️ Confirmar Exclusão?", ["Não", "Sim"], horizontal=True, key="confirmar_excluir")
    if confirmar == "Sim" and st.button("🗑️ Confirmar Exclusão"):
        id_excluir = linha_x.get("ID", "")
        if not id_excluir:
            st.error("❌ Este registro não possui ID. Não é possível excluir com segurança.")
        else:
            ok = delete_row_by_id("timesheet.csv", id_excluir)
            if ok:
                st.success("✅ Registro excluído com sucesso!")
                st.rerun()

    st.markdown("---")
    st.subheader("📥 Exportar Dados")
    df_export = df_visual.copy()
    buffer = df_export.to_csv(index=False, sep=CSV_SEP, encoding=CSV_ENC).encode()
    st.download_button(label="📥 Baixar Tabela", data=buffer, file_name="timesheet_filtrado.csv", mime="text/csv")

# =============================================
# CONTEÚDO: PERFORMANCE — IA
# =============================================

elif menu == "📊 Avaliação de Performance — IA":
    st.title("📊 Avaliação de Performance com IA")
    usuario_logado = st.session_state.username

    df_ts, _ = carregar_base("timesheet.csv")
    df_ts = normalizar_coluna_horas(df_ts)
    df_ts["Data"] = pd.to_datetime(df_ts["Data"], errors="coerce")

    if df_ts.empty:
        st.info("⚠️ Não há dados no timesheet para avaliar.")
        st.stop()

    if usuario_logado not in ADMIN_USERS:
        st.error("🚫 Você não tem permissão para acessar a Avaliação de Performance.")
        st.stop()

    st.markdown("### 🤖 Gerando relatório com IA")

    lista_projetos = sorted(df_ts["Projeto"].dropna().unique().tolist())
    projeto_escolhido = st.multiselect("Selecione o Projeto para análise:", ["Todos os Projetos"] + lista_projetos)

    lista_colaboradores = sorted(df_ts["Nome"].dropna().unique().tolist())
    colaborador_escolhido = st.multiselect("Selecione o Colaborador para análise:", ["Todos os Colaboradores"] + lista_colaboradores)

    df_ts["Ano"] = df_ts["Data"].dt.year
    df_ts["Mes"] = df_ts["Data"].dt.strftime('%m - %B')
    anos_disponiveis = sorted(df_ts["Ano"].dropna().unique().tolist())
    ano_escolhido = st.multiselect("Selecione o Ano:", ["Todos os Anos"] + anos_disponiveis)

    meses_disponiveis = df_ts["Mes"].dropna().unique().tolist()
    meses_disponiveis_ordenados = sorted(meses_disponiveis, key=lambda x: int(str(x).split(" - ")[0]))
    mes_escolhido = st.multiselect("Selecione o Mês:", ["Todos os Meses"] + meses_disponiveis_ordenados)

    df_f = df_ts.copy()
    if "Todos os Projetos" not in projeto_escolhido:
        df_f = df_f[df_f["Projeto"].isin(projeto_escolhido)]
    if "Todos os Colaboradores" not in colaborador_escolhido:
        df_f = df_f[df_f["Nome"].isin(colaborador_escolhido)]
    if "Todos os Anos" not in ano_escolhido:
        df_f = df_f[df_f["Ano"].isin(ano_escolhido)]
    if "Todos os Meses" not in mes_escolhido:
        df_f = df_f[df_f["Mes"].isin(mes_escolhido)]

    if df_f.empty:
        st.info("⚠️ Nenhum registro encontrado para o filtro selecionado.")
        st.stop()

    client = OpenAI(api_key=st.secrets["openai"]["api_key"])

    dados_markdown = df_f.fillna("").astype(str).to_markdown(index=False)
    prompt = f"""
    Você é um consultor especialista em gestão de tempo, produtividade e análise de performance.

    Analise os dados do timesheet abaixo e gere um relatório completo e estruturado contendo:
    - ✅ Resumo executivo
    - ✅ Principais indicadores
    - ✅ Gargalos e desvios
    - ✅ Recomendações de melhorias operacionais
    - ✅ Conclusões finais

    Seja objetivo, técnico e claro. Utilize contagens, percentuais e análises de tendência.

    ### Dados do Timesheet:
    {dados_markdown}
    """

    if st.button("🚀 Gerar Relatório de Performance"):
        with st.spinner("A IA está gerando o relatório..."):
            resposta = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Você é um especialista em análise de produtividade corporativa e performance."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )
            texto_relatorio = resposta.choices[0].message.content
            st.success("✅ Relatório gerado com sucesso!")
            st.markdown("### 📄 Relatório Gerado:")
            st.markdown(texto_relatorio)

            doc = Document()
            style = doc.styles["Normal"]
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            doc.add_heading("📊 Relatório de Avaliação de Performance", level=1)
            doc.add_paragraph(f"Projetos: {', '.join(projeto_escolhido) if projeto_escolhido else '—'}")
            doc.add_paragraph(f"Colaboradores: {', '.join(colaborador_escolhido) if colaborador_escolhido else '—'}")
            doc.add_paragraph(f"Data da geração: {datetime.today().strftime('%Y-%m-%d')}")
            doc.add_paragraph("\n")
            for linha in (texto_relatorio or "").split("\n"):
                if linha.strip().startswith("#"):
                    nivel = linha.count("#")
                    texto = linha.replace("#", "").strip()
                    doc.add_heading(texto, level=min(nivel, 4))
                else:
                    doc.add_paragraph(linha.strip())
            buff = BytesIO()
            doc.save(buff)
            buff.seek(0)
            st.download_button(
                label="📥 Baixar Relatório em Word",
                data=buff,
                file_name="relatorio_performance.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
